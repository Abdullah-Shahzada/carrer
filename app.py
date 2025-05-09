import streamlit as st
from groq import Groq
from docx import Document

# Initialize Groq client
client = Groq(api_key="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")

# Function to create a .doc file
def create_doc_file(user_name, career_recommendations, skill_gap_analysis, learning_resources, job_websites, success_stories, psychological_analysis, developer_name):
    doc = Document()
    doc.add_heading(f"Career Recommendation Report for {user_name}", 0)
    doc.add_paragraph(f"Dear {user_name},\n")
    doc.add_paragraph("Based on the information you provided, we have analyzed your profile and here are your career recommendations:\n")

    doc.add_heading("Career Recommendations", level=1)
    doc.add_paragraph(career_recommendations)

    doc.add_heading("Skill Gap Analysis", level=1)
    doc.add_paragraph(skill_gap_analysis)

    doc.add_heading("Certifications and Learning Resources", level=1)
    doc.add_paragraph(learning_resources)

    doc.add_heading("Suggested Job Posting Websites", level=1)
    for website in job_websites:
        doc.add_paragraph(website, style="ListBullet")

    doc.add_heading("Success Stories", level=1)
    doc.add_paragraph(success_stories)

    doc.add_heading("Psychological Factor Analysis", level=1)
    doc.add_paragraph(psychological_analysis)

    # Closing note
    doc.add_paragraph(f"\n\nThanks {user_name} for trusting our AI Career Advisory System! 🚀")
    doc.add_paragraph(f"Developed by: {developer_name}")

    # Save the doc
    doc.save("career_recommendations.docx")

# Streamlit App Title
st.title("AI-Personalized learning 🚀")
st.markdown("**Developed by: Career_Advisor**")

# Sidebar Navigation
st.sidebar.title("Navigation")
section = st.sidebar.radio("Go to", ["User Input", "Career Recommendations", "Skill Gap Analysis", "Certifications & Resources", "Success Stories", "Psychological Analysis", "Download DOC"])
with st.sidebar.expander("Job Platforms"):
    st.markdown("[LinkedIn](https://www.linkedin.com/jobs/)")
    st.markdown("[Indeed](https://www.indeed.com/)")
    st.markdown("[Glassdoor](https://www.glassdoor.com/Job/index.htm)")
    st.markdown("[Monster](https://www.monster.com/)")
    st.markdown("[AngelList (Startups)](https://angel.co/jobs)")
    st.markdown("[Remote.co (Remote Jobs)](https://remote.co/remote-jobs/)")

# Initialize session states
for key in ["user_name", "career_recommendations", "skill_gap_analysis", "learning_resources", "psychological_analysis", "success_stories"]:
    if key not in st.session_state:
        st.session_state[key] = ""

# --- USER INPUT SECTION ---
if section == "User Input":
    st.header("Enter Your Details")
    user_name = st.text_input("Please enter your full name:")
    if user_name:
        st.session_state["user_name"] = user_name
    skills = st.text_input("Enter your skills (comma-separated):")
    interests = st.text_input("Enter your interests (comma-separated):")
    career_goals = st.text_input("Enter your career goals:")
    education = st.text_input("Enter your educational background:")
    industries = st.text_input("Enter industries of interest (comma-separated):")
    experience_level = st.selectbox("Select your experience level:", ["Beginner", "Intermediate", "Expert"])
    psychological_factors = st.selectbox("How well do you handle stress and challenges?", ["Very Well", "Moderate", "Not Well"])
    stress_handling = st.selectbox("How do you handle work pressure?", ["Thrives under pressure", "Handles pressure well", "Struggles with pressure"])
    problem_solving = st.selectbox("How do you approach complex problems?", ["Enjoy solving challenges", "Can solve problems with guidance", "Finds problem-solving stressful"])

    if st.button("Generate Reports"):
        if skills and interests and career_goals and education and industries:
            user_input = f"""
            User Name: {user_name}
            Skills: {skills}
            Interests: {interests}
            Career Goals: {career_goals}
            Educational Background: {education}
            Industries of Interest: {industries}
            Experience Level: {experience_level}
            Psychological Factors: {psychological_factors}
            Stress Handling: {stress_handling}
            Problem Solving: {problem_solving}
            """

            with st.spinner("Generating Career Recommendations..."):
                completion = client.chat.completions.create(
                    model="llama-3.1-8b-instant",
                    messages=[{
                        "role": "user",
                        "content": f"Suggest 3 career options based on:\n{user_input}"
                    }],
                    temperature=0.7,
                    max_tokens=1024,
                )
                st.session_state["career_recommendations"] = completion.choices[0].message.content

            with st.spinner("Generating Skill Gap Analysis..."):
                completion = client.chat.completions.create(
                    model="llama-3.1-8b-instant",
                    messages=[{
                        "role": "user",
                        "content": f"Analyze skill gaps based on:\n{user_input}"
                    }],
                    temperature=0.7,
                    max_tokens=1024,
                )
                st.session_state["skill_gap_analysis"] = completion.choices[0].message.content

            with st.spinner("Generating Certifications and Resources..."):
                completion = client.chat.completions.create(
                    model="llama-3.1-8b-instant",
                    messages=[{
                        "role": "user",
                        "content": f"Suggest certifications, courses, and books for:\n{user_input}"
                    }],
                    temperature=0.7,
                    max_tokens=1024,
                )
                st.session_state["learning_resources"] = completion.choices[0].message.content

            with st.spinner("Generating Psychological Factor Analysis..."):
                completion = client.chat.completions.create(
                    model="llama-3.1-8b-instant",
                    messages=[{
                        "role": "user",
                        "content": f"Analyze psychological factors and suggest career paths for:\n{user_input}"
                    }],
                    temperature=0.7,
                    max_tokens=1024,
                )
                st.session_state["psychological_analysis"] = completion.choices[0].message.content

            with st.spinner("Fetching Success Stories..."):
                completion = client.chat.completions.create(
                    model="llama-3.1-8b-instant",
                    messages=[{
                        "role": "user",
                        "content": f"Provide real-world success stories of famous people who succeeded in the field(s) of {industries}. Mention their names, achievements, roadmap they follow and a short inspiring note for each."
                    }],
                    temperature=0.7,
                    max_tokens=1024,
                )
                st.session_state["success_stories"] = completion.choices[0].message.content

            st.success("All sections generated! 🎉 Now explore from the sidebar.")

        else:
            st.error("Please fill all fields.")

# --- CAREER RECOMMENDATIONS SECTION ---
if section == "Career Recommendations":
    st.header("Career Recommendations")
    if st.session_state["career_recommendations"]:
        st.markdown(st.session_state["career_recommendations"])

        st.subheader("Chat About Career Recommendations")
        if "career_chat" not in st.session_state:
            st.session_state["career_chat"] = []

        career_question = st.text_input("Ask a question about career recommendations:")
        if st.button("Send (Career Chat)"):
            if career_question:
                with st.spinner("Thinking..."):
                    response = client.chat.completions.create(
                        model="llama-3.1-8b-instant",
                        messages=[
                            {"role": "assistant", "content": st.session_state["career_recommendations"]},
                            {"role": "user", "content": career_question}
                        ],
                        temperature=0.7,
                        max_tokens=1024,
                    )
                    answer = response.choices[0].message.content
                    st.session_state["career_chat"].append((career_question, answer))
            else:
                st.error("Please enter a question.")

        for q, a in st.session_state["career_chat"]:
            st.markdown(f"**You:** {q}")
            st.markdown(f"**AI:** {a}")
    else:
        st.info("Generate career recommendations first.")

# --- SKILL GAP ANALYSIS SECTION ---
if section == "Skill Gap Analysis":
    st.header("Skill Gap Analysis")
    if st.session_state["skill_gap_analysis"]:
        st.markdown(st.session_state["skill_gap_analysis"])

        st.subheader("Chat About Skill Gaps")
        if "skill_chat" not in st.session_state:
            st.session_state["skill_chat"] = []

        skill_question = st.text_input("Ask a question about skill gaps:")
        if st.button("Send (Skill Chat)"):
            if skill_question:
                with st.spinner("Thinking..."):
                    response = client.chat.completions.create(
                        model="llama-3.1-8b-instant",
                        messages=[
                            {"role": "assistant", "content": st.session_state["skill_gap_analysis"]},
                            {"role": "user", "content": skill_question}
                        ],
                        temperature=0.7,
                        max_tokens=1024,
                    )
                    answer = response.choices[0].message.content
                    st.session_state["skill_chat"].append((skill_question, answer))
            else:
                st.error("Please enter a question.")

        for q, a in st.session_state["skill_chat"]:
            st.markdown(f"**You:** {q}")
            st.markdown(f"**AI:** {a}")
    else:
        st.info("Generate skill gap analysis first.")

# --- CERTIFICATIONS AND LEARNING RESOURCES SECTION ---
if section == "Certifications & Resources":
    st.header("Certifications and Learning Resources")
    if st.session_state["learning_resources"]:
        st.markdown(st.session_state["learning_resources"])

        st.subheader("Chat About Learning Resources")
        if "learning_chat" not in st.session_state:
            st.session_state["learning_chat"] = []

        learning_question = st.text_input("Ask a question about learning resources:")
        if st.button("Send (Learning Chat)"):
            if learning_question:
                with st.spinner("Thinking..."):
                    response = client.chat.completions.create(
                        model="llama-3.1-8b-instant",
                        messages=[
                            {"role": "assistant", "content": st.session_state["learning_resources"]},
                            {"role": "user", "content": learning_question}
                        ],
                        temperature=0.7,
                        max_tokens=1024,
                    )
                    answer = response.choices[0].message.content
                    st.session_state["learning_chat"].append((learning_question, answer))
            else:
                st.error("Please enter a question.")

        for q, a in st.session_state["learning_chat"]:
            st.markdown(f"**You:** {q}")
            st.markdown(f"**AI:** {a}")
    else:
        st.info("Generate learning resources first.")
# --- SUCCESS STORIES SECTION ---
if section == "Success Stories":
    st.header("Success Stories")
    if st.session_state["success_stories"]:
        st.markdown(st.session_state["success_stories"])

        st.subheader("Chat About Success Stories")
        if "success_chat" not in st.session_state:
            st.session_state["success_chat"] = []

        success_question = st.text_input("Ask a question about success stories:")
        if st.button("Send (Success Chat)"):
            if success_question:
                with st.spinner("Thinking..."):
                    response = client.chat.completions.create(
                        model="llama-3.1-8b-instant",
                        messages=[
                            {"role": "assistant", "content": st.session_state["success_stories"]},
                            {"role": "user", "content": success_question}
                        ],
                        temperature=0.7,
                        max_tokens=1024,
                    )
                    answer = response.choices[0].message.content
                    st.session_state["success_chat"].append((success_question, answer))
            else:
                st.error("Please enter a question.")

        for q, a in st.session_state["success_chat"]:
            st.markdown(f"**You:** {q}")
            st.markdown(f"**AI:** {a}")
    else:
        st.info("Generate success stories first.")

# --- PSYCHOLOGICAL FACTOR ANALYSIS SECTION ---
if section == "Psychological Analysis":
    st.header("Psychological Factor Analysis")
    if st.session_state["psychological_analysis"]:
        st.markdown(st.session_state["psychological_analysis"])

        st.subheader("Chat About Psychological Analysis")
        if "psych_chat" not in st.session_state:
            st.session_state["psych_chat"] = []

        psych_question = st.text_input("Ask a question about psychological analysis:")
        if st.button("Send (Psych Chat)"):
            if psych_question:
                with st.spinner("Thinking..."):
                    response = client.chat.completions.create(
                        model="llama-3.1-8b-instant",
                        messages=[
                            {"role": "assistant", "content": st.session_state["psychological_analysis"]},
                            {"role": "user", "content": psych_question}
                        ],
                        temperature=0.7,
                        max_tokens=1024,
                    )
                    answer = response.choices[0].message.content
                    st.session_state["psych_chat"].append((psych_question, answer))
            else:
                st.error("Please enter a question.")

        for q, a in st.session_state["psych_chat"]:
            st.markdown(f"**You:** {q}")
            st.markdown(f"**AI:** {a}")
    else:
        st.info("Generate psychological analysis first.")

# --- DOWNLOAD DOC SECTION ---
if section == "Download DOC":
    st.header("Save and Download Recommendations")
    if all([st.session_state["career_recommendations"], st.session_state["skill_gap_analysis"], st.session_state["learning_resources"], st.session_state["psychological_analysis"], st.session_state["success_stories"]]):
        developer_name = "Carreer_Advisor"
        create_doc_file(
            st.session_state.get("user_name", "User"),
            st.session_state["career_recommendations"],
            st.session_state["skill_gap_analysis"],
            st.session_state["learning_resources"],
            [
                "LinkedIn", "Indeed", "Glassdoor", "Monster", "AngelList", "Dice", "Remote.co"
            ],
            st.session_state["success_stories"],
            st.session_state["psychological_analysis"],
            developer_name
        )
        with open("career_recommendations.docx", "rb") as file:
            st.download_button(
                label="Download Career Report DOC",
                data=file,
                file_name=f"{st.session_state.get('user_name', 'User')}_career_recommendations.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        st.info("Please generate all sections first.")

